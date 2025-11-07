import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import logging
from pathlib import Path
from datetime import datetime
import threading
import time
import subprocess
import sys

logger = logging.getLogger("DiceSensei.MainWindow")

class MainWindow:
    def __init__(self, root, ollama_manager):
        self.root = root
        self.ollama_manager = ollama_manager
        self.current_document = None
        self.document_content = ""
        self.chat_history = []
        self.is_processing = False
        
        self.setup_ui()
        self.load_config()
    
    def setup_ui(self):
        """Configura la interfaz de usuario principal"""
        self.setup_styles()
        self.create_main_frame()
        self.create_header()
        self.create_chat_area()
        self.create_quick_actions()
        self.create_status_bar()
        self.create_progress_bar()
        
        self.apply_theme()

    def create_progress_bar(self):
        """Crea la barra de progreso (oculta inicialmente)"""
        self.progress_frame = ttk.Frame(self.main_frame)
        self.progress_frame.grid(row=5, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        self.progress_frame.grid_remove() 
        
        self.progress_label = tk.Label(
            self.progress_frame,
            text="",
            font=('Arial', 9),
            fg=self.light_theme['accent']
        )
        self.progress_label.grid(row=0, column=0, sticky=tk.W)
        
        self.progress_bar = ttk.Progressbar(
            self.progress_frame,
            mode='indeterminate',
            length=300
        )
        self.progress_bar.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=5)
        
        self.progress_frame.columnconfigure(0, weight=1)
    
    def show_progress(self, message="Procesando..."):
        """Muestra la barra de progreso"""
        self.progress_label.config(text=message)
        self.progress_frame.grid()
        self.progress_bar.start()
        self.root.update()
    
    def hide_progress(self):
        """Oculta la barra de progreso"""
        self.progress_frame.grid_remove()
        self.progress_bar.stop()
        self.root.update()
    
    def update_progress_text(self, message):
        """Actualiza el texto de la barra de progreso"""
        self.progress_label.config(text=message)
        self.root.update()
    
    def setup_styles(self):
        """Configura los estilos de la aplicación"""
        self.style = ttk.Style()
        
        self.style.theme_use('clam')
        
        self.light_theme = {
            'bg': '#f8f9fa',
            'fg': '#212529',
            'accent': '#007bff',
            'secondary': '#6c757d',
            'success': '#28a745',
            'warning': '#ffc107',
            'danger': '#dc3545',
            'chat_user_bg': '#007bff',
            'chat_user_fg': 'white',
            'chat_bot_bg': '#e9ecef',
            'chat_bot_fg': '#212529'
        }
        
        self.dark_theme = {
            'bg': '#212529',
            'fg': '#f8f9fa',
            'accent': '#0d6efd',
            'secondary': '#adb5bd',
            'success': '#198754',
            'warning': '#ffca2c',
            'danger': '#dc3545',
            'chat_user_bg': '#0d6efd',
            'chat_user_fg': 'white',
            'chat_bot_bg': '#495057',
            'chat_bot_fg': '#f8f9fa'
        }
    
    def create_main_frame(self):
        """Crea el frame principal"""
        self.main_frame = ttk.Frame(self.root, padding="10")
        self.main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        self.main_frame.columnconfigure(1, weight=1)
        self.main_frame.rowconfigure(2, weight=1)
    
    def create_header(self):
        """Crea la cabecera de la aplicación"""
        header_frame = ttk.Frame(self.main_frame)
        header_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        
        title_frame = ttk.Frame(header_frame)
        title_frame.grid(row=0, column=0, sticky=tk.W)
        
        self.title_label = tk.Label(
            title_frame,
            text="🎲 DiceSensei",
            font=('Arial', 18, 'bold'),
            fg=self.light_theme['accent']
        )
        self.title_label.grid(row=0, column=0, padx=(0, 10))
        
        self.subtitle_label = tk.Label(
            title_frame,
            text="Tu asistente de estudio inteligente",
            font=('Arial', 10),
            fg=self.light_theme['secondary']
        )
        self.subtitle_label.grid(row=1, column=0)
        
        action_frame = ttk.Frame(header_frame)
        action_frame.grid(row=0, column=1, sticky=tk.E)
        
        self.load_btn = tk.Button(
            action_frame,
            text="📚 Cargar Documento",
            command=self.load_document,
            bg=self.light_theme['accent'],
            fg='white',
            font=('Arial', 10, 'bold'),
            padx=15,
            pady=5,
            relief='flat',
            bd=0
        )
        self.load_btn.grid(row=0, column=0, padx=5)
        
        self.settings_btn = tk.Button(
            action_frame,
            text="⚙️ Configuración",
            command=self.open_settings,
            bg=self.light_theme['secondary'],
            fg='white',
            font=('Arial', 10),
            padx=15,
            pady=5,
            relief='flat',
            bd=0
        )
        self.settings_btn.grid(row=0, column=1, padx=5)
        
        header_frame.columnconfigure(0, weight=1)
    
    def create_chat_area(self):
        """Crea el área de chat"""
        chat_frame = ttk.LabelFrame(
            self.main_frame, 
            text="Chat de Estudio",
            padding="10"
        )
        chat_frame.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=10)
        
        self.chat_history_text = scrolledtext.ScrolledText(
            chat_frame,
            height=20,
            width=70,
            font=('Arial', 10),
            bg='white',
            fg='black',
            wrap=tk.WORD,
            state=tk.DISABLED,
            padx=10,
            pady=5
        )
        self.chat_history_text.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        
        self.setup_chat_tags()
        
        input_frame = ttk.Frame(chat_frame)
        input_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E))
        
        self.input_entry = tk.Entry(
            input_frame,
            font=('Arial', 11),
            bg='white',
            fg='black'
        )
        self.input_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 10))
        self.input_entry.bind('<Return>', lambda e: self.send_message())
        
        self.send_btn = tk.Button(
            input_frame,
            text="📤 Enviar",
            command=self.send_message,
            bg=self.light_theme['success'],
            fg='white',
            font=('Arial', 10, 'bold'),
            padx=20
        )
        self.send_btn.grid(row=0, column=1)
        
        input_frame.columnconfigure(0, weight=1)
        chat_frame.columnconfigure(0, weight=1)
        chat_frame.rowconfigure(0, weight=1)
    
    def setup_chat_tags(self):
        """Configura los estilos para diferentes tipos de mensajes"""
        self.chat_history_text.tag_configure(
            "user",
            background=self.light_theme['chat_user_bg'],
            foreground=self.light_theme['chat_user_fg'],
            justify='right'
        )
        
        self.chat_history_text.tag_configure(
            "bot",
            background=self.light_theme['chat_bot_bg'],
            foreground=self.light_theme['chat_bot_fg'],
            justify='left'
        )
        
        self.chat_history_text.tag_configure(
            "system",
            foreground=self.light_theme['secondary'],
            font=('Arial', 9, 'italic')
        )
        
        # Tag para mensaje "pensando"
        self.chat_history_text.tag_configure(
            "thinking",
            foreground=self.light_theme['secondary'],
            font=('Arial', 10, 'italic')
        )
    
    def create_quick_actions(self):
        """Crea los botones de acciones rápidas"""
        actions_frame = ttk.LabelFrame(
            self.main_frame,
            text="Preguntas Rápidas",
            padding="5"
        )
        actions_frame.grid(row=3, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=10)
        
        quick_actions = [
            ("📖 Resumir documento", self.quick_summarize),
            ("❓ Preguntas clave", self.quick_questions),
            ("🔍 Conceptos principales", self.quick_concepts),
            ("💡 Ideas importantes", self.quick_ideas),
            ("📝 Explicar detalladamente", self.quick_explain)
        ]
        
        for i, (text, command) in enumerate(quick_actions):
            btn = tk.Button(
                actions_frame,
                text=text,
                command=command,
                bg=self.light_theme['secondary'],
                fg='white',
                font=('Arial', 9),
                padx=10,
                pady=5,
                relief='flat'
            )
            btn.grid(row=0, column=i, padx=2)
    
    def create_status_bar(self):
        """Crea la barra de estado"""
        self.status_frame = ttk.Frame(self.main_frame)
        self.status_frame.grid(row=4, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(10, 0))
        
        self.status_label = tk.Label(
            self.status_frame,
            text="Listo",
            font=('Arial', 9),
            fg=self.light_theme['secondary']
        )
        self.status_label.grid(row=0, column=0, sticky=tk.W)

        self.hardware_label = tk.Label(
        self.status_frame,
        text="",  
        font=('Arial', 8),
        fg=self.light_theme['secondary']
        )
        self.hardware_label.grid(row=0, column=1)
        
        self.document_label = tk.Label(
            self.status_frame,
            text="Documento: Ninguno",
            font=('Arial', 9),
            fg=self.light_theme['secondary']
        )
        self.document_label.grid(row=0, column=2, sticky=tk.E)
        
        self.status_frame.columnconfigure(0, weight=1)
        self.status_frame.columnconfigure(1, weight=0)
        self.status_frame.columnconfigure(2, weight=1)

    def update_hardware_info(self, cores, ram):
        """Actualiza la información de hardware en la barra de estado"""
        self.hardware_label.config(text=f"🖥️ {cores} cores | 💾 {ram}GB")
    
    def apply_theme(self):
        """Aplica el tema actual a la interfaz"""
        pass
    
    def load_config(self):
        """Carga la configuración"""
        pass
    
    def load_document(self):
        """Carga un documento para estudiar"""
        file_types = [
            ("Documentos", "*.txt *.pdf *.docx *.md *.rtf"),
            ("Textos", "*.txt"),
            ("PDFs", "*.pdf"),
            ("Word", "*.docx"),
            ("Markdown", "*.md"),
            ("Todos", "*.*")
        ]
        
        file_path = filedialog.askopenfilename(
            title="Seleccionar documento de estudio",
            filetypes=file_types
        )
        
        if file_path:
            self.process_document(file_path)
    
    def install_missing_dependency(self, package_name):
        """Instala una dependencia faltante automáticamente"""
        try:
            self.update_progress_text(f"📦 Instalando {package_name}...")
            subprocess.check_call([
                sys.executable, "-m", "pip", "install", package_name
            ], timeout=60)
            self.update_progress_text(f"✅ {package_name} instalado correctamente")
            time.sleep(1)
            return True
        except Exception as e:
            self.update_progress_text(f"❌ Error instalando {package_name}")
            logger.error(f"Error instalando {package_name}: {e}")
            return False
    
    def process_document(self, file_path):
        """Procesa el documento cargado con manejo de diferentes formatos"""
        try:
            self.current_document = file_path
            doc_name = Path(file_path).name
            
            self.show_progress(f"Cargando {doc_name}...")
            
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.pdf':
                self.document_content = self._read_pdf(file_path)
            elif file_extension == '.docx':
                self.document_content = self._read_docx(file_path)
            elif file_extension in ['.txt', '.md', '.rtf']:
                self.document_content = self._read_text(file_path)
            else:
                raise ValueError(f"Formato no soportado: {file_extension}")
            
            self.hide_progress()
            
            self.document_label.config(text=f"Documento: {doc_name}")
            self.add_to_chat("system", f"📚 Documento cargado: {doc_name}")
            self.add_to_chat("system", f"📄 Tamaño: {len(self.document_content)} caracteres")
            self.add_to_chat("system", "¡Listo! Puedes hacer preguntas sobre este material.")
            
            self.update_status(f"Documento '{doc_name}' cargado correctamente")
            
        except Exception as e:
            self.hide_progress()
            messagebox.showerror("Error", f"No se pudo cargar el documento:\n{str(e)}")
            logger.error(f"Error cargando documento: {e}")
    
    def _read_pdf(self, file_path):
        """Lee un archivo PDF con instalación automática de dependencias"""
        try:
            self.update_progress_text("Extrayendo texto del PDF...")
            
            try:
                import PyPDF2
            except ImportError:
                if not self.install_missing_dependency("pypdf2"):
                    raise Exception("No se pudo instalar PyPDF2 para leer PDFs")
                import PyPDF2
            
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    total_pages = len(pdf_reader.pages)
                    
                    for page_num in range(total_pages):
                        self.update_progress_text(f"Procesando página {page_num + 1}/{total_pages}...")
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
                    
                    if not text.strip():
                        raise Exception("El PDF parece estar vacío o ser una imagen")
                    
                    return text.strip()
                    
            except Exception as e:
                logger.warning(f"PyPDF2 falló: {e}")
                
                try:
                    try:
                        from pypdf import PdfReader
                    except ImportError:
                        if not self.install_missing_dependency("pypdf"):
                            raise Exception("No se pudo instalar pypdf para leer PDFs")
                        from pypdf import PdfReader
                    
                    with open(file_path, 'rb') as file:
                        pdf_reader = PdfReader(file)
                        text = ""
                        total_pages = len(pdf_reader.pages)
                        
                        for page_num, page in enumerate(pdf_reader.pages):
                            self.update_progress_text(f"Procesando página {page_num + 1}/{total_pages}...")
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n\n"
                        
                        if not text.strip():
                            raise Exception("El PDF no contiene texto extraíble")
                        
                        return text.strip()
                        
                except Exception as e2:
                    logger.error(f"pypdf también falló: {e2}")
                    raise Exception(f"No se pudo extraer texto del PDF. El archivo puede estar dañado o ser una imagen.")
                    
        except Exception as e:
            raise Exception(f"Error procesando PDF: {str(e)}")
    
    def _read_docx(self, file_path):
        """Lee un archivo DOCX con instalación automática de dependencias"""
        try:
            self.update_progress_text("Extrayendo texto del documento Word...")
            
            try:
                import docx
            except ImportError:
                if not self.install_missing_dependency("python-docx"):
                    raise Exception("No se pudo instalar python-docx para leer documentos Word")
                import docx
            
            doc = docx.Document(file_path)
            text = ""
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise Exception("El documento Word parece estar vacío")
            
            return text.strip()
            
        except Exception as e:
            raise Exception(f"Error leyendo documento Word: {str(e)}")
    
    def _read_text(self, file_path):
        """Lee archivos de texto con detección de codificación"""
        self.update_progress_text("Leyendo archivo de texto...")
        
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    logger.info(f"✅ Archivo leído con codificación: {encoding}")
                    return content
            except UnicodeDecodeError:
                continue
            except Exception as e:
                continue
        
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
                return content.decode('utf-8', errors='ignore')
        except Exception as e:
            raise Exception(f"No se pudo leer el archivo con ninguna codificación: {str(e)}")
    
    def send_message(self):
        """Envía un mensaje del usuario"""
        if self.is_processing:
            return
            
        message = self.input_entry.get().strip()
        if not message:
            return
        
        self.input_entry.delete(0, tk.END)
        self.add_to_chat("user", message)
        
        threading.Thread(
            target=self.process_user_message,
            args=(message,),
            daemon=True
        ).start()
    
    def process_user_message(self, message):
        """Procesa el mensaje del usuario y genera respuesta con feedback visual"""
        try:
            self.is_processing = True
            self.send_btn.config(state=tk.DISABLED)
            self.input_entry.config(state=tk.DISABLED)
            
            thinking_id = self.add_thinking_message()
            
            context = self.document_content if self.document_content else None
            
            def generate_response():
                try:
                    response = self.ollama_manager.generate_response(
                        prompt=message,
                        context=context
                    )
                    self.root.after(0, lambda: self.show_response(thinking_id, response))
                except Exception as e:
                    error_msg = f"Error: No se pudo generar respuesta ({str(e)})"
                    self.root.after(0, lambda: self.show_response(thinking_id, error_msg))
            
            threading.Thread(target=generate_response, daemon=True).start()
            
        except Exception as e:
            self.is_processing = False
            self.send_btn.config(state=tk.NORMAL)
            self.input_entry.config(state=tk.NORMAL)
            error_msg = f"Error: No se pudo procesar tu mensaje ({str(e)})"
            self.add_to_chat("system", error_msg)
            self.update_status("Error")
            logger.error(f"Error procesando mensaje: {e}")
    
    def add_thinking_message(self):
        """Añade un mensaje de 'pensando' y retorna su ID"""
        self.chat_history_text.config(state=tk.NORMAL)
        
        self.chat_history_text.insert(tk.END, "🎲 DiceSensei está analizando...", "thinking")
        self.chat_history_text.see(tk.END)
        self.chat_history_text.config(state=tk.DISABLED)
        
        return self.chat_history_text.index(tk.END)
    
    def show_response(self, thinking_id, response):
        """Reemplaza el mensaje de 'pensando' con la respuesta real"""
        self.chat_history_text.config(state=tk.NORMAL)
        
        self.chat_history_text.delete(f"{thinking_id}-1c", thinking_id)
        
        self.chat_history_text.insert(tk.END, f"{response}\n\n", "bot")
        
        self.chat_history_text.see(tk.END)
        self.chat_history_text.config(state=tk.DISABLED)
        
        self.is_processing = False
        self.send_btn.config(state=tk.NORMAL)
        self.input_entry.config(state=tk.NORMAL)
        self.update_status("Listo")
    
    def add_to_chat(self, sender, message):
        """Añade un mensaje al historial de chat"""
        self.chat_history_text.config(state=tk.NORMAL)
        
        timestamp = datetime.now().strftime("%H:%M")
        
        if sender == "user":
            prefix = "👤 Tú"
            tag = "user"
        elif sender == "bot":
            prefix = "🎲 DiceSensei"
            tag = "bot"
        else:
            prefix = "⚙️ Sistema"
            tag = "system"
        
        self.chat_history_text.insert(tk.END, f"[{timestamp}] {prefix}: ", tag)
        self.chat_history_text.insert(tk.END, f"{message}\n\n", tag)
        
        self.chat_history_text.see(tk.END)
        self.chat_history_text.config(state=tk.DISABLED)
    
    def update_status(self, message):
        """Actualiza la barra de estado"""
        self.status_label.config(text=message)
        self.root.update()
    
    def quick_summarize(self):
        """Acción rápida: resumir documento"""
        if not self.document_content:
            messagebox.showwarning("Advertencia", "Primero carga un documento")
            return
        
        self.input_entry.delete(0, tk.END)
        self.input_entry.insert(0, "Por favor, haz un resumen detallado de este documento")
        self.send_message()
    
    def quick_questions(self):
        """Acción rápida: generar preguntas clave"""
        if not self.document_content:
            messagebox.showwarning("Advertencia", "Primero carga un documento")
            return
        
        self.input_entry.delete(0, tk.END)
        self.input_entry.insert(0, "Genera 5 preguntas clave para evaluar la comprensión de este material")
        self.send_message()
    
    def quick_concepts(self):
        """Acción rápida: explicar conceptos principales"""
        if not self.document_content:
            messagebox.showwarning("Advertencia", "Primero carga un documento")
            return
        
        self.input_entry.delete(0, tk.END)
        self.input_entry.insert(0, "Explica los conceptos principales del documento")
        self.send_message()
    
    def quick_ideas(self):
        """Acción rápida: ideas importantes"""
        if not self.document_content:
            messagebox.showwarning("Advertencia", "Primero carga un documento")
            return
        
        self.input_entry.delete(0, tk.END)
        self.input_entry.insert(0, "¿Cuáles son las ideas más importantes de este documento?")
        self.send_message()
    
    def quick_explain(self):
        """Acción rápida: explicar detalladamente"""
        if not self.document_content:
            messagebox.showwarning("Advertencia", "Primero carga un documento")
            return
        
        self.input_entry.delete(0, tk.END)
        self.input_entry.insert(0, "Explica este documento como si me lo estuvieras enseñando")
        self.send_message()
    
    def open_settings(self):
        """Abre la ventana de configuración"""
        messagebox.showinfo("Configuración", "Ventana de configuración en desarrollo")

    def on_closing(self):
        """Maneja el cierre de la aplicación"""
        try:
            self.ollama_manager.stop_server()
        except:
            pass
        self.root.destroy()
